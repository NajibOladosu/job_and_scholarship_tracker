"""
Document parsing service for extracting text from various document formats.
"""
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import logging
import os

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Service class for parsing documents and extracting text.
    """

    def parse_document(self, file_path, document_type):
        """
        Parse document and extract text based on file type.

        Args:
            file_path (str): Path to the document file
            document_type (str): Type hint (resume, transcript, certificate, other)

        Returns:
            dict: Dictionary containing:
                - success (bool): Whether parsing succeeded
                - text (str): Extracted text content
                - error (str): Error message if failed
        """
        try:
            # Determine file type from extension
            _, extension = os.path.splitext(file_path)
            extension = extension.lower()

            logger.info(f"Parsing document: {file_path} (type: {document_type}, ext: {extension})")

            # Route to appropriate parser
            if extension == '.pdf':
                result = self._parse_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                result = self._parse_docx(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                result = self._parse_image(file_path)
            elif extension == '.txt':
                result = self._parse_txt(file_path)
            else:
                result = {
                    'success': False,
                    'text': '',
                    'error': f'Unsupported file type: {extension}'
                }

            if result['success']:
                logger.info(f"Successfully parsed document. Extracted {len(result['text'])} characters.")

            return result

        except Exception as e:
            logger.error(f"Unexpected error parsing document {file_path}: {e}")
            return {
                'success': False,
                'text': '',
                'error': str(e)
            }

    def _parse_pdf(self, file_path):
        """
        Extract text from PDF file using pdfplumber.

        Args:
            file_path (str): Path to PDF file

        Returns:
            dict: Parsing result
        """
        try:
            text_parts = []

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            full_text = '\n'.join(text_parts)

            if not full_text.strip():
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text could be extracted from PDF'
                }

            return {
                'success': True,
                'text': full_text,
                'error': ''
            }

        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return {
                'success': False,
                'text': '',
                'error': f'PDF parsing error: {str(e)}'
            }

    def _parse_docx(self, file_path):
        """
        Extract text from DOCX file using python-docx.

        Args:
            file_path (str): Path to DOCX file

        Returns:
            dict: Parsing result
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = '\n'.join(text_parts)

            if not full_text.strip():
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text could be extracted from DOCX'
                }

            return {
                'success': True,
                'text': full_text,
                'error': ''
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return {
                'success': False,
                'text': '',
                'error': f'DOCX parsing error: {str(e)}'
            }

    def _parse_image(self, file_path):
        """
        Extract text from image using OCR (pytesseract).

        Args:
            file_path (str): Path to image file

        Returns:
            dict: Parsing result
        """
        try:
            # Open image
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image)

            if not text.strip():
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text could be extracted from image (OCR found nothing)'
                }

            return {
                'success': True,
                'text': text,
                'error': ''
            }

        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract OCR not installed or not in PATH")
            return {
                'success': False,
                'text': '',
                'error': 'Tesseract OCR not installed. OCR functionality unavailable.'
            }
        except Exception as e:
            logger.error(f"Error parsing image {file_path}: {e}")
            return {
                'success': False,
                'text': '',
                'error': f'Image OCR error: {str(e)}'
            }

    def _parse_txt(self, file_path):
        """
        Read text from TXT file.

        Args:
            file_path (str): Path to TXT file

        Returns:
            dict: Parsing result
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            return {
                'success': True,
                'text': text,
                'error': ''
            }

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return {
                    'success': True,
                    'text': text,
                    'error': ''
                }
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {e}")
                return {
                    'success': False,
                    'text': '',
                    'error': f'Text file encoding error: {str(e)}'
                }
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return {
                'success': False,
                'text': '',
                'error': f'Text file reading error: {str(e)}'
            }


# Singleton instance
_document_parser = None


def get_document_parser():
    """
    Get or create singleton instance of DocumentParser.

    Returns:
        DocumentParser: The singleton service instance
    """
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser
